HOME = 'C:\\Temp\\'
FILENAME = 'new_doc.docx'
HEADER = 'This is original document'
TEXT = """String 1
String 2
String 3"""
            
from docx import Document
import os.path

# Функция создана для удобства, чтобы не создавать и не удалять файл каждый раз
def create_doc(path, text):
    if os.path.isfile(path):
        # Если файл создан, удаляем его
        os.remove(path)
    # Если файла нет или он удален, создаем новый с известным соержимым
    doc = Document()
    header = doc.add_heading(HEADER, 0)
    doc.add_paragraph(TEXT).add_run
    doc.save(path)
    return 1

def replace_string(path, search_string, new_string):
    doc = Document(path)
    for i in range(len(doc.paragraphs)):    
        text = doc.paragraphs[i].text
        doc.paragraphs[i].text = text.replace(search_string, new_string)
    doc.save(path)
    
path = HOME + FILENAME
r = create_doc(path, TEXT)
print(r)

replace_string (path, 'original', 'changed')
replace_string (path, 'String 2', 'One more string')